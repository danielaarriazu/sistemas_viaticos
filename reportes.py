import os
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import viaticos


def _img(img_path, width, height):
    """Retorna Image si el archivo existe, sino un Spacer del mismo tamaño."""
    if img_path and os.path.exists(img_path):
        return Image(img_path, width=width, height=height)
    return Spacer(width, height)

def obtener_nombre_unico(nombre_base, extension=".docx"):
    i = 1
    nombre_final = f"{nombre_base}{extension}"
    while os.path.exists(nombre_final):
        nombre_final = f"{nombre_base}_{i}{extension}"
        i += 1
    return nombre_final
    

def build_receipt_docx(path: str, data: dict) -> None:
    """Genera exclusivamente el recibo en formato Word (.docx) con el Anexo 38."""
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    doc = docx.Document()
    
    # Configuración formal de márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Tipografía base (Arial 11)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Encabezado estructurado sin bordes
    table_hdr = doc.add_table(rows=1, cols=2)
    table_hdr.autofit = True
    table_hdr.cell(0, 0).paragraphs[0].add_run("ARMADA ARGENTINA").bold = True
    
    p_hdr_right = table_hdr.cell(0, 1).paragraphs[0]
    p_hdr_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_hdr_right.add_run("ANEXO 38\n(703.e)").bold = True

    doc.add_paragraph() # Espaciador

    # Título Oficial
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("RECIBO POR VIÁTICOS POR COMISIONES DEL SERVICIO\nPERSONAL MILITAR")
    run_title.bold = True
    run_title.font.size = Pt(12)

    doc.add_paragraph()

    # Tratamiento de datos financieros y texto del cuerpo
    total_float  = data.get("total_calculado", 0.0)
    total_txt    = viaticos.format_money(total_float)
    total_letras = viaticos.importe_a_letras(total_float)
    pct_aplicado = data.get("porcentaje_aplicado", 100.0)
    texto_pct_nota = f" (al {pct_aplicado}%)" if pct_aplicado != 100.0 else ""
    
    texto_cuerpo = (
        f"--------- Recibí de la JEFATURA ADMINISTRATIVA FINANCIERA DE LA ARMADA "
        f"la suma de {total_letras} ({total_txt}), correspondiente a "
        f"{data.get('detalle_breakdown', '')} del viático diario de "
        f"{viaticos.format_money(data.get('importe_diario', 0.0))} – "
        f"({data.get('categoria', '')}){texto_pct_nota}. " 
        f"Desde {data.get('lugar_origen', '')} el día {data.get('fecha_salida', '')} a las {data.get('hora_salida', '')} "
        f"hasta {data.get('destino_comision', '')} el día {data.get('fecha_regreso', '')} a las {data.get('hora_regreso', '')}, "
        f"con motivo de {data.get('motivo', '')}. Por orden {data.get('orden', '')}."
    )
    p_body = doc.add_paragraph(texto_cuerpo)
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body.paragraph_format.line_spacing = 1.25

    doc.add_paragraph()

    # Fecha de Emisión (Derecha)
    def fecha_es_local(fecha_str):
        if not fecha_str: return ""
        dia, mes, anio = fecha_str.split("/")
        meses = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
                 "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        return f"{dia} de {meses[int(mes)]} de {anio}"

    lugar_emision = data.get("lugar_emision", "Buenos Aires")
    fecha_fmt = fecha_es_local(data.get("fecha_emision_recibo", ""))
    p_emision = doc.add_paragraph(f"{lugar_emision}, {fecha_fmt}.")
    p_emision.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Cuadro de firmas del Causante e imágenes de validación
    causante_txt = (
        f"Firma: .................................................\n"
        f"Aclaración: {data.get('nombre', '')} {data.get('apellido', '')}\n"
        f"M.R. {data.get('mr', '')} – {data.get('jerarquia_desc', '')}\n"
        f"DNI: {data.get('dni', '')}\n"
        f"COD. Alfa/Numérico: {data.get('cod_destino', '')}\n"
        f"Destino de Revista: {data.get('destino_revista_desc', '')}"
    )
    p_causante = doc.add_paragraph(causante_txt)
    p_causante.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Inserción de Sellos y Firmas del Detall y Director
    table_signs = doc.add_table(rows=1, cols=2)
    cell_detall = table_signs.cell(0, 0)
    if data.get("sello_detall_path") and os.path.exists(data.get("sello_detall_path")):
        cell_detall.paragraphs[0].add_run().add_picture(data["sello_detall_path"], width=Cm(4.0))
    if data.get("firma_detall_path") and os.path.exists(data.get("firma_detall_path")):
        cell_detall.add_paragraph().add_run().add_picture(data["firma_detall_path"], width=Cm(4.0))

    cell_dir = table_signs.cell(0, 1)
    if data.get("sello_director_path") and os.path.exists(data.get("sello_director_path")):
        cell_dir.paragraphs[0].add_run().add_picture(data["sello_director_path"], width=Cm(4.0))
    if data.get("firma_director_path") and os.path.exists(data.get("firma_director_path")):
        cell_dir.add_paragraph().add_run().add_picture(data["firma_director_path"], width=Cm(4.0))

    doc.save(path)